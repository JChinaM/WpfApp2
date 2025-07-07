import os
import sys
import requests
from bs4 import BeautifulSoup, NavigableString, Tag
from urllib.parse import urljoin, urlparse
from docx import Document
from docx.shared import Inches

if len(sys.argv) < 3:
    print("Использование: python script.py <URL> <Путь_для_сохранения_файла.docx>")
    sys.exit(1)

url = sys.argv[1]
output_docx_file = sys.argv[2]

output_folder = os.path.dirname(os.path.abspath(output_docx_file))
images_folder = os.path.join(output_folder, "images_osa")
os.makedirs(images_folder, exist_ok=True)

headers = {"User-Agent": "Mozilla/5.0"}

print(f"[+] Загружаю страницу: {url}")
response = requests.get(url, headers=headers)
if response.status_code != 200:
    print(f"[!] Ошибка загрузки страницы: {response.status_code}")
    sys.exit(1)

soup = BeautifulSoup(response.content, "html.parser")

def extract_text(selector_list):
    for sel in selector_list:
        el = soup.select_one(sel)
        if el and el.get_text(strip=True):
            return el.get_text(strip=True)
    return None

title = extract_text(["h1.entry-title", "h1", "title"]) or "Без названия"
date = extract_text(["time", ".date", ".posted-on"]) or "Без даты"
author = "Без автора"

content = soup.select_one("div.content_item.posts_item")
if not content:
    print("[!] Не удалось найти блок с текстом статьи!")
    sys.exit(1)

doc = Document()
doc.add_heading(title, 0)
doc.add_paragraph(f"Дата: {date}")
doc.add_paragraph(f"Автор: {author}")
doc.add_paragraph("\n")

image_counter = 1

def process_node(node):
    global image_counter
    if isinstance(node, NavigableString):
        text = node.strip()
        if text:
            doc.add_paragraph(text)
    elif isinstance(node, Tag):
        if node.name in ["p", "div", "span"]:
            for child in node.contents:
                process_node(child)
        elif node.name in ["h1", "h2", "h3", "h4"]:
            heading = node.get_text(strip=True)
            if heading:
                doc.add_heading(heading, level=2)
        elif node.name in ["ul", "ol"]:
            for li in node.find_all("li", recursive=False):
                li_text = li.get_text(strip=True)
                doc.add_paragraph(f"- {li_text}")
        elif node.name == "blockquote":
            quote = node.get_text(strip=True)
            if quote:
                doc.add_paragraph(f'"{quote}"', style='Intense Quote')
        elif node.name == "img":
            img_src = node.get("src")
            if img_src:
                full_img_url = urljoin(url, img_src)
                parsed = urlparse(img_src)
                ext = os.path.splitext(parsed.path)[1] or ".jpg"
                img_filename = f"image{image_counter}{ext}"
                img_path = os.path.join(images_folder, img_filename)

                try:
                    img_data = requests.get(full_img_url, headers=headers, timeout=10).content
                    with open(img_path, "wb") as f:
                        f.write(img_data)
                    print(f"[+] Сохранено изображение: {img_filename}")
                    doc.add_picture(img_path, width=Inches(5))
                    image_counter += 1
                except Exception as e:
                    print(f"[!] Ошибка при загрузке изображения: {e}")
        else:
            for child in node.contents:
                process_node(child)

process_node(content)

doc.save(output_docx_file)
print(f"\n[??] Статья полностью сохранена в: {output_docx_file}")
print(f"[??] Картинки сохранены в папке: {images_folder}")
